from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openai
import time

app = Flask(__name__)
CORS(app)

app.config["UPLOAD_FOLDER"] = "uploads"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

openai.api_key = os.getenv("OPENAI_API_KEY")

def match_terms_to_paragraph(paragraph, terms_df):
    paragraph_lower = paragraph.lower()
    matched = []
    for _, row in terms_df.iterrows():
        term = row['Term'].strip().lower()
        if term in paragraph_lower:
            matched.append(f"{row['Term']} = {row['Translation']}")
    return matched

@app.route("/translate", methods=["POST"])
def translate_file():

    # 🧹 حذف ملفات إكسل ووورد أقدم من ٤٨ ساعة
    now = time.time()
    for f in os.listdir(app.config["UPLOAD_FOLDER"]):
        full_path = os.path.join(app.config["UPLOAD_FOLDER"], f)
        if not os.path.isfile(full_path):
            continue
        age = now - os.path.getmtime(full_path)
        if (f.endswith(".xlsx") or f.endswith(".docx")) and age > 172800:
            os.remove(full_path)

    # 🔍 التحقق من وجود ملف مؤقت مطابق
    import difflib

    def normalize(text):
        return text.lower().strip().replace("\n", "").replace("  ", " ")

    def is_similar(a, b):
        return difflib.SequenceMatcher(None, normalize(a), normalize(b)).ratio() >= 0.8

    matching_temp_file = None
    for f in os.listdir(app.config["UPLOAD_FOLDER"]):
        if "_temp_translated.xlsx" in f.lower() and filename_clean.lower() in f.lower():
            matching_temp_file = os.path.join(app.config["UPLOAD_FOLDER"], f)
            break

    if matching_temp_file:
        try:
            temp_df = pd.read_excel(matching_temp_file)
            if "Extracted Paragraphs" in temp_df.columns:
                temp_paragraphs = temp_df["Extracted Paragraphs"].dropna().tolist()[:3]
                if file_ext == ".csv":
                    new_paragraphs = pd.read_csv(input_path)["Extracted Paragraphs"].dropna().tolist()[:3]
                else:
                    doc = Document(input_path)
                    new_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:3]
                matches = sum(is_similar(tp, np) for tp, np in zip(temp_paragraphs, new_paragraphs))
                if matches >= 2:
                    paragraphs_df = temp_df.copy()
                    excel_output = matching_temp_file.replace("_temp_translated.xlsx", f"_{time.strftime('%Y%m%d_%H%M')}_translated.xlsx")
                    os.rename(matching_temp_file, excel_output)
                    return send_file(excel_output, as_attachment=True)
        except Exception as e:
            print(f"[تحذير] فشل التحقق من الملف المؤقت: {e}")

    service = request.form.get("service", "Translation_in_Excel")

    uploaded_file = request.files.get("file")
    glossary_file = request.files.get("glossary")

    source_lang = request.form.get("source_lang", "English")
    target_lang = request.form.get("target_lang", "Arabic")

    if not uploaded_file:
        return jsonify({"error": "لم يتم رفع أي ملف."}), 400

    filename_raw = uploaded_file.filename
    filename_clean = os.path.splitext(secure_filename(filename_raw))[0]
    file_ext = os.path.splitext(filename_raw)[1].lower()
    input_path = os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(filename_raw))
    uploaded_file.save(input_path)

    time.sleep(0.2)

    glossary_path = None

    if service == "Translation_in_Excel":
        try:
            if file_ext == ".csv":
                paragraphs_df = pd.read_csv(input_path)
                if "Extracted Paragraphs" not in paragraphs_df.columns:
                    return jsonify({"error": "الملف لا يحتوي على عمود 'Extracted Paragraphs'."}), 400
            elif file_ext == ".docx":
                doc = Document(input_path)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                paragraphs_df = pd.DataFrame(paragraphs, columns=["Extracted Paragraphs"])
            else:
                return jsonify({"error": "صيغة الملف غير مدعومة. الرجاء رفع ملف .docx أو .csv"}), 400
        except Exception as e:
            return jsonify({"error": f"فشل في قراءة الملف: {str(e)}"}), 400

        glossary_df = pd.DataFrame(columns=["Term", "Translation"])
        if glossary_file:
            try:
                glossary_name = secure_filename(glossary_file.filename)
                glossary_path = os.path.join(app.config["UPLOAD_FOLDER"], glossary_name)
                glossary_file.save(glossary_path)
                doc = Document(glossary_path)
                if not doc.tables:
                    return jsonify({"error": "ملف المصطلحات لا يحتوي على أي جدول."}), 400
                table = doc.tables[0]
                if len(table.columns) < 2:
                    return jsonify({"error": "الجدول في ملف المصطلحات يجب أن يحتوي على عمودين على الأقل."}), 400
                terms = [row.cells[0].text.strip() for row in table.rows[1:]]
                translations_glossary = [row.cells[1].text.strip() for row in table.rows[1:]]
                glossary_df = pd.DataFrame({"Term": terms, "Translation": translations_glossary})
            except Exception as e:
                return jsonify({"error": f"فشل في قراءة ملف المصطلحات: {str(e)}"}), 400

        excel_output = os.path.join(app.config["UPLOAD_FOLDER"], f"{filename_clean}_translated.xlsx")

        if os.path.exists(excel_output):
            paragraphs_df_existing = pd.read_excel(excel_output)
            if "Translation" in paragraphs_df_existing.columns:
                paragraphs_df["Translation"] = paragraphs_df_existing["Translation"]
                paragraphs_df["Status"] = paragraphs_df_existing.get("Status", "")
            else:
                paragraphs_df["Translation"] = ""
                paragraphs_df["Status"] = ""
        else:
            paragraphs_df["Translation"] = ""
            paragraphs_df["Status"] = ""

        intro_paragraphs = paragraphs_df.iloc[:2, 0].tolist()
        joined_intro = "\n".join(intro_paragraphs)
        context_prompt = f"""You are given the beginning of a technical or regulatory document.
Your task is to generate a single clear English sentence that describes the main topic or context of the document.

Content:
{joined_intro}

Context hint:"""

        try:
            context_response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert technical summarizer."},
                    {"role": "user", "content": context_prompt}
                ],
                temperature=0.3,
                max_tokens=50
            )
            context_hint = context_response.choices[0].message.content.strip()
        except Exception as e:
            context_hint = "[تعذر توليد سياق المستند]"
            print(f"[تحذير] فشل في توليد السياق: {e}")

        Translation_Prompt = f"""As a Translator, your task is to convert the following text from {source_lang} to {target_lang},
ensuring that the terminology is precise and the language flows naturally.
Each word and phrase must align with the context of this field, as the content pertains to legal and technical domains.
Your focus should be on producing a text that is clear, accurate, and easy to read
while fully preserving the original meaning and structure.
The translation must be suitable for a reverse translation that closely matches the original,
avoiding any extra content, interpretation, or logical gaps. Do not include translator notes.

Document Context:
{context_hint}
"""

        for idx, row in paragraphs_df.iterrows():
            if row["Translation"] and row["Translation"] != "missing translation":
                continue

            paragraph = row["Extracted Paragraphs"]
            
            matched_terms = match_terms_to_paragraph(paragraph, glossary_df)
            glossary_section = ""
            if matched_terms:
                glossary_text = "\n".join(matched_terms)
                glossary_section = f"""Glossary:
{glossary_text}

"""

            prompt = f"""{Translation_Prompt}
{glossary_section}
{paragraph}"""

            try:
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are a professional English-to-Arabic legal and technical translator."},
                        {"role": "user", "content": prompt.strip()}
                    ],
                    temperature=0.0,
                    max_tokens=1000
                )
                paragraphs_df.at[idx, "Translation"] = response.choices[0].message.content.strip()
                paragraphs_df.at[idx, "Status"] = "Success"
            except Exception as e:
                paragraphs_df.at[idx, "Translation"] = "missing translation"
                paragraphs_df.at[idx, "Status"] = "Failed"

        paragraphs_df.to_excel(excel_output, index=False)

        if os.path.exists(input_path):
            os.remove(input_path)
        if glossary_path and os.path.exists(glossary_path):
            os.remove(glossary_path)

        response = send_file(excel_output, as_attachment=True)

        if os.path.exists(excel_output):
            os.remove(excel_output)

        return response

    elif service == "Convert_to_Word":
        if file_ext != ".xlsx":
            return jsonify({"error": "يرجى رفع ملف Excel صالح يحتوي على الترجمة."}), 400
        try:
            df = pd.read_excel(input_path)
        except Exception as e:
            return jsonify({"error": f"فشل في قراءة ملف Excel: {str(e)}"}), 400

        if "Translation" not in df.columns:
            return jsonify({"error": "ملف الترجمة لا يحتوي على عمود 'Translation'."}), 400

        word_output = os.path.join(app.config["UPLOAD_FOLDER"], f"{filename_clean}_translated.docx")
        doc = Document()

        # تحديد التنسيق بناءً على اللغة المستهدفة
        is_arabic = target_lang.lower().strip() in ["arabic", "ar"]

        style = doc.styles['Normal']
        if is_arabic:
            style.font.name = 'Simplified Arabic'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Simplified Arabic')
            style.font.size = Pt(14)
            doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            doc.styles['Normal'].paragraph_format.left_to_right = False
            doc.styles['Normal'].paragraph_format.right_to_left = True
            doc.settings.language_id = 1025  # Arabic - Saudi Arabia
        else:
            style.font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            style.font.size = Pt(13)
            doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.styles['Normal'].paragraph_format.left_to_right = True
            doc.styles['Normal'].paragraph_format.right_to_left = False
            doc.settings.language_id = 1033  # English - US

        style = doc.styles['Normal']
        style.font.name = 'Simplified Arabic'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Simplified Arabic')
        style.font.size = Pt(14)
        section = doc.sections[0]
        section.right_margin = section.left_margin
        doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        for text in df["Translation"].fillna(""):
            para = doc.add_paragraph(str(text).strip())
            para.paragraph_format.space_after = Pt(0)

        doc.save(word_output)

        if os.path.exists(input_path):
            os.remove(input_path)

        return send_file(word_output, as_attachment=True)

    else:
        return jsonify({"error": "نوع الخدمة غير معروف. يرجى التحقق من قيمة 'service'."}), 400

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
