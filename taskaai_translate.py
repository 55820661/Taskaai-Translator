import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from dotenv import load_dotenv
import openai
import shutil

print("📦 Loading environment variables...")
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

source_lang = "English"
target_lang = "Arabic"

output_excel_path = "translated_output_ready.xlsx"
output_word_path = "final_translated_output.docx"
temp_excel_path = "translated_output_temp.xlsx"
glossary_docx_path = "translated_terms.docx"

possible_paths = ["extracted_paragraphs.csv", "extracted_paragraphs.docx"]
paragraphs_csv_path = None
for path in possible_paths:
    if os.path.exists(path):
        paragraphs_csv_path = path
        print(f"📥 Found input file: {path}")
        break
if paragraphs_csv_path is None:
    print("❌ No input file found.")
    exit(1)

file_ext = os.path.splitext(paragraphs_csv_path)[1].lower()
if file_ext == ".csv":
    paragraphs_df = pd.read_csv(paragraphs_csv_path)
elif file_ext == ".docx":
    doc = Document(paragraphs_csv_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    paragraphs_df = pd.DataFrame(paragraphs, columns=["Extracted Paragraphs"])
else:
    print("❌ Unsupported file type.")
    exit(1)

if os.path.exists(temp_excel_path):
    print("🔁 Resuming from temporary translation file...")
    temp_df = pd.read_excel(temp_excel_path)
    temp_translations = temp_df.get("Translation", [""] * len(paragraphs_df)).tolist()
    translations = temp_translations + ["" for _ in range(len(paragraphs_df) - len(temp_translations))]
else:
    translations = ["" for _ in range(len(paragraphs_df))]

if not os.path.exists(glossary_docx_path):
    print("⚠️ لم يتم العثور على ملف المصطلحات.")
    glossary_df = pd.DataFrame(columns=["Term", "Translation"])
else:
    doc = Document(glossary_docx_path)
    table = doc.tables[0]
    terms = [row.cells[0].text.strip() for row in table.rows[1:]]
    translations_glossary = [row.cells[1].text.strip() for row in table.rows[1:]]
    glossary_df = pd.DataFrame({"Term": terms, "Translation": translations_glossary})

print("🧠 Generating document context from first two paragraphs...")
intro_paragraphs = paragraphs_df.iloc[:2, 0].tolist()
joined_intro = "\n".join(intro_paragraphs)
context_prompt = f"""You are given the beginning of a technical or regulatory document.
Your task is to generate a single clear {source_lang} sentence that describes the main topic or context of the document.

Content:
{joined_intro}

Context hint:"""

context_response = openai.ChatCompletion.create(
    model="gpt-4",
    messages=[
        {"role": "system", "content": "You are an expert technical summarizer."},
        {"role": "user", "content": context_prompt}
    ],
    temperature=0.3,
    max_tokens=50
)
context_hint = context_response.choices[0].message.content.strip()
print(f"📘 Context hint generated: {context_hint}")

def match_terms_to_paragraph(paragraph, terms_df):
    paragraph_lower = paragraph.lower()
    matched = []
    for _, row in terms_df.iterrows():
        term = row['Term'].strip().lower()
        if term in paragraph_lower:
            matched.append(f"{row['Term']} = {row['Translation']}")
    return matched

print("📝 Starting translation of paragraphs...")
for i, paragraph in enumerate(paragraphs_df.iloc[:, 0]):
    if isinstance(translations[i], str) and translations[i].strip():
        print(f"⏩ Skipping already translated paragraph {i+1}")
        continue

    print(f"🔁 Processing paragraph {i+1}/{len(paragraphs_df)}...")
    if '----media/' in paragraph:
        translations[i] = paragraph.strip()
        continue

    matched_terms = match_terms_to_paragraph(paragraph, glossary_df)
    glossary_text = "\n".join(matched_terms) if matched_terms else "[No relevant glossary terms]"
    prompt = f"""[STRICT HAMADA TRANSLATION PROMPT]

Document Context: {context_hint}

Translate the following text from {source_lang} to {target_lang} using precise and literal translation.
Do not paraphrase or summarize. Use the glossary below exactly as provided if matching terms are found.
Maintain the original sentence structure and order. Avoid any interpretation or stylistic changes.

Glossary:
{glossary_text}

Text:
{paragraph}"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"You are a professional {source_lang}-to-{target_lang} legal and technical translator."},
                {"role": "user", "content": prompt.strip()}
            ],
            temperature=0.0,
            max_tokens=1000
        )
        translations[i] = response.choices[0].message.content.strip()
    except Exception as e:
        translations[i] = f"[Error] {str(e)}"

    pd.DataFrame({
        'Extracted Paragraphs': paragraphs_df.iloc[:len(translations), 0],
        'Translation': translations
    }).to_excel(temp_excel_path, index=False)

if len([t for t in translations if isinstance(t, str) and t.strip()]) < len(paragraphs_df):
    print("⚠️ Translation incomplete. Not all paragraphs have been translated. Resume script to continue.")
    exit(0)

print("💾 Saving translated content to Excel...")
pd.DataFrame({
    'Extracted Paragraphs': paragraphs_df.iloc[:, 0],
    'Translation': translations
}).to_excel(output_excel_path, index=False)

# حذف الملف المؤقت بعد الانتهاء
if os.path.exists(temp_excel_path):
    os.remove(temp_excel_path)

print("📄 Creating Word file with translations...")
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Simplified Arabic'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Simplified Arabic')
style.font.size = Pt(14)
section = doc.sections[0]
section.right_margin = section.left_margin
doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

for text in translations:
    if isinstance(text, str):
        para = doc.add_paragraph(text.strip())
        para.paragraph_format.space_after = Pt(0)

doc.save(output_word_path)
print(f"✅ Translation complete. Word file saved: {output_word_path}")
